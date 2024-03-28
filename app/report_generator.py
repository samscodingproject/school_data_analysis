from docx import Document
from docx.shared import Inches
import docx.shared
import io
import pandas as pd
import datetime
import docx.oxml
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement 

def set_table_borders(table):
    tbl = table._tbl
    tbl.attrib[qn('w:tblBorders')] = ''
    for cell in tbl.iter_tcs():
        tcPr = cell.tcPr
        tcBorders = OxmlElement('w:tcBorders')
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '4')
        top.set(qn('w:space'), '0')
        top.set(qn('w:color'), 'auto')
        tcBorders.append(top)

        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '4')
        left.set(qn('w:space'), '0')
        left.set(qn('w:color'), 'auto')
        tcBorders.append(left)

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '4')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')
        tcBorders.append(bottom)

        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'single')
        right.set(qn('w:sz'), '4')
        right.set(qn('w:space'), '0')
        right.set(qn('w:color'), 'auto')
        tcBorders.append(right)

        tcPr.append(tcBorders)

def generate_student_report(marks_df, student_ids, attendance_df=None, overall_attendance=None):
    document = Document()
    print(f"Number of student IDs: {len(student_ids)}")  # Debugging statement

    # Get the current year
    current_year = datetime.datetime.now().year

    

    # Dictionary to map long subject names to their shortened versions
    subject_name_mapping = {
        'Community & Family Studies': 'CAFS',
        'Earth & Environmental Science': 'Earth & Env. Sci.',
        'English Extension 1': 'Eng Ext 1',
        'English Standard': 'Eng Std',
        'Information Processes & Technology': 'IPT',
        'Mathematics Advanced': 'Math Adv',
        'Mathematics Extension 1': 'Math Ext 1',
        'Mathematics Standard 2': 'Math Std 2',
        'Software Design & Development': 'SDD'
    }

    for student_id in student_ids:
        print(f"Generating report for student ID: {student_id}")  # Debugging statement
        student_marks_data = marks_df[marks_df['StudentID'] == student_id]
        print(f"Number of subjects for student {student_id}: {len(student_marks_data)}")  # Debugging statement

        # Check if 'Student Name' exists and is not empty, otherwise use 'StudentID'
        if 'Student Name' in student_marks_data.columns and not pd.isnull(student_marks_data['Student Name'].values[0]):
            student_identifier = student_marks_data['Student Name'].values[0]
        else:
            student_identifier = str(student_id)

        # Create a table with one row and two cells
        header_table = document.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.allow_autofit = False
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(6)

        # Add the school logo to the left cell
        left_cell = header_table.cell(0, 0)
        left_cell.width = Inches(1.5)
        para = left_cell.paragraphs[0]
        run = para.add_run()
        run.add_picture('/Users/Sam/Desktop/New schools data analysis/app/static/images/st-mary-logo.png', width=Inches(1.2))

        # Add the school name and program name to the right cell
        right_cell = header_table.cell(0, 1)
        para = right_cell.paragraphs[0]
        run = para.add_run('St Marys Senior High School\n')
        run.bold = True
        run.font.size = docx.shared.Pt(14)
        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        para = right_cell.add_paragraph('ACADEMIC IMPROVEMENT PROGRAM')
        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

        # Add the student's identifier, date, and school year
        if attendance_df is not None:
            school_year = attendance_df[attendance_df['StudentID'] == student_id]['School Year'].values[0]
        
        else:
            school_year = 'N/A'
        student_info = f"{student_identifier} - {datetime.datetime.now().strftime('%d %B %Y')} - Year {school_year}"
        document.add_heading(student_info, level=1)

        # Academic Results Section
        document.add_heading(f'Academic Results {current_year}:', level=2)

        # Check if any student has Term 4 data
        has_term_4_data = False
        for _, row in student_marks_data.iterrows():
            if 'T4' in row and not pd.isnull(row['T4']):
                has_term_4_data = True
                break

        # Determine the number of columns based on Term 4 data availability
        num_columns = 7 if has_term_4_data else 6
        academic_results_table = document.add_table(rows=1, cols=num_columns)

        # Set column widths
        academic_results_table.columns[0].width = Inches(2.1)  # Subject column
        for i in range(1, num_columns):
            academic_results_table.columns[i].width = Inches(1.0)  # Other columns

        hdr_cells = academic_results_table.rows[0].cells
        hdr_cells[0].text = 'Subject'
        hdr_cells[1].text = 'Term 1 Mark'
        hdr_cells[2].text = 'Term 2 Mark'
        hdr_cells[3].text = 'Term 3 Mark'
        if has_term_4_data:
            hdr_cells[4].text = 'Term 4 Mark'
            hdr_cells[5].text = 'Final Mark'
            hdr_cells[6].text = 'Final z-Score'
        else:
            hdr_cells[4].text = 'Final Mark'
            hdr_cells[5].text = 'Final z-Score'

        for _, row in student_marks_data.iterrows():
            row_cells = academic_results_table.add_row().cells
            subject = row['Subject']
            if subject in subject_name_mapping:
                subject = subject_name_mapping[subject]
            row_cells[0].text = subject
            row_cells[1].text = f"{row['T1']}%" if not pd.isnull(row['T1']) else '-'
            row_cells[2].text = f"{row['T2']}%" if not pd.isnull(row['T2']) else '-'
            row_cells[3].text = f"{row['T3']}%" if not pd.isnull(row['T3']) else '-'
            if has_term_4_data:
                row_cells[4].text = f"{row['T4']}%" if 'T4' in row and not pd.isnull(row['T4']) else '-'
                row_cells[5].text = f"{row['CalculatedFinalMark']}%"
                row_cells[6].text = str(row['zScore'])
            else:
                row_cells[4].text = f"{row['CalculatedFinalMark']}%"
                row_cells[5].text = str(row['zScore'])

        set_table_borders(academic_results_table)

        # Attendance Section
        document.add_heading(f'Attendance {current_year}:', level=2)
        attendance_table = document.add_table(rows=1, cols=2)

        hdr_cells = attendance_table.rows[0].cells
        hdr_cells[0].text = 'Overall Attendance'
        hdr_cells[1].text = 'Percentage'

        if overall_attendance is not None:
            student_overall_attendance = overall_attendance[overall_attendance['StudentID'] == student_id]['AttendancePercentage'].values[0]
            overall_attendance_row = attendance_table.add_row().cells
            overall_attendance_row[0].text = 'Overall Attendance'
            overall_attendance_row[1].text = f"{str(student_overall_attendance)}%"

        if attendance_df is not None:
            student_attendance_data = attendance_df[attendance_df['StudentID'] == student_id]
            for _, row in student_attendance_data.iterrows():
                subject = row['Subject']
                attendance_percentage = row['AttendancePercentage']
                new_row = attendance_table.add_row().cells
                new_row[0].text = subject
                new_row[1].text = f"{str(attendance_percentage)}%"

        set_table_borders(attendance_table)

        # Deputy Principal Sign Off
        document.add_paragraph(" ")
        document.add_paragraph("Deputy Principal Sign Off: ___________________________")

        # Add a page break before the next student's report, if there are more students
        if student_id != student_ids[-1]:
            document.add_page_break()

    print("Report generation completed")  # Debugging statement
    # Save the document to a BytesIO object and return it
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)
    return file_stream